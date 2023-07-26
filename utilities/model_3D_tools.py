# 29012023
# 3D model utility Functions
# Python 3


"""Tools for 3D model management..

List of tools:
        * models_3D - object with list of 3D models
        * Can get different formats of 3D models files.
        * get_obj_label_from_file_name(file_name = ''): - get object label from file name



Caution:

Example:

"""


# Imports

import os
from os import walk, path
import math
from datetime import datetime


from docx import Document
from docx.shared import Inches
from docx.shared import Pt
import numpy as np
import pandas as pd   # https://pandas.pydata.org/
import dir_tools
import csv
import dir_tools

def get_obj_label_from_file_name(file_name = ''):
    """
    get_obj_label_from_file_name(file_name = ''): - get object label from file name
    Input arguments:
        - file_name - File name string

    Output
        - obj_label - object label string
       ..

    Caution:

    Example:

    Notes:
        Update as needed.

    """

    dir_tools.check_dir_name_str(file_name)
    FileName, extension = path.splitext(file_name)
    obj_label = FileName[(FileName.find('_label_')+len('_label_')):]

    return obj_label



#########################################################################################################################################################################################
#########################################################################################################################################################################################
#########################################################################################################################################################################################




class models_3D_list:

    # ROOT_DIR = os.path.normpath(os.getcwd())
    ROOT_DIR = os.path.dirname(os.path.abspath(
        "top_level_file.txt"))  # Get project root https://www.adamsmith.haus/python/answers/how-to-get-the-path-of-the-root-project-structure-in-python

    def __init__(self, model3D_extension=None, model3D_dir= ROOT_DIR, limit=math.inf, out_model3D_list = ROOT_DIR):
        """
        :param model3D_extension: tuple of strings of file extensions which need to search in the model3D_dir. by default
          model3D_extension = ['.STL', '.stl']

        :param model3D_dir: Folder path name where files will be searched. If no directory provided the root project durectory is used as the default.

        :param limit: Limit on number of models per model list file. default = infinity.

        :param out_model3D_list: Output directory path for lists of 3D models. If no directory provided the root project durectory is used as the default.

        :
        """
        # TODO: extend to other CAD file types. We should add here function that will convert from other model formats into STL
        #       Here we shiuld add funcrion that will automaticaly detect model file extension and if needed will convert it to STL.
        dir_tools.check_path_name_str(model3D_dir)
        dir_tools.check_path_name_str(out_model3D_list)


        assert path.isdir(model3D_dir)   #https://www.geeksforgeeks.org/python-os-path-isdir-method/  # Check if we have direcory to scan for 3D models.
        now = datetime.now()  # Get time
        self.time = now.strftime("%d%m%Y_%H%M%S")
        out_model3D_list = os.path.normpath(os.path.abspath(os.path.join(out_model3D_list, 'data_list_' + self.time)))
        out_model3D_list = dir_tools.create_path(out_model3D_list)
        self.out_model3D_list = out_model3D_list
        self.list_3Dmodels_txt = 'list_3Dmodels_' + self.time  # name of the file that includes list of the models in txt format
        self.list_3Dmodels_docx = 'list_3Dmodels_' + self.time # name of the file that includes list of the models in doc format
        self.list_3Dmodels_csv = 'list_3Dmodels_' + self.time # name of the file that includes list of the models in doc format


        self.limit = limit
        self.model3D_dir = model3D_dir
        if model3D_extension is None:
            model3D_extension = ['.STL', '.stl']
        self.model3D_extension = model3D_extension

        self.path_list_txt =''
        self.path_list_doc =''
        self.path_list_csv =''

        self.generate_3Dmodel_list()



    def generate_3Dmodel_list(self):
        list_3Dmodels_txt_count = 1
        # TXT
        self.path_list_txt = path.normpath(path.abspath(os.path.join(self.out_model3D_list, self.list_3Dmodels_txt + '_' + str(list_3Dmodels_txt_count) + '.txt')))
        list_file_txt = open( self.path_list_txt , 'w')  # txt file

        # Doc
        list_file_doc = Document()  # https://python-docx.readthedocs.io/en/latest/index.html
        list_file_doc.add_heading('List of 3D Models' + "\n", 0)
        FontSizeMain = 14
        DocxListObj = list_file_doc.add_paragraph()
        DocxListObj.add_run("\n" + self.time + "\n").font.size = Pt(FontSizeMain)

        # CSV
        self.path_list_csv = path.normpath(path.abspath(os.path.join(self.out_model3D_list, self.list_3Dmodels_csv + '_' + str(list_3Dmodels_txt_count) + '.csv')))
        list_file_csv = open( self.path_list_csv , 'w', newline='')  # csv file  https://www.pythontutorial.net/python-basics/python-write-csv-file/
         # https://stackoverflow.com/questions/3348460/csv-file-written-with-python-has-blank-lines-between-each-row
        writer = csv.writer(list_file_csv)
        header = ['index', 'full_path' , 'path_no_file_name', 'file_name', 'label']
        # write the header
        writer.writerow(header)

        count = 0
        for r, d, files in walk(self.model3D_dir):  #https://www.tutorialspoint.com/python/os_walk.htm
            for file in files:
                _, extension = path.splitext(file)
                if extension in self.model3D_extension:              # Write 3D model file to list
                    count += 1
                    # txt
                    path2write = path.normpath(path.abspath(path.join(r, file))) + '\n'
                    list_file_txt.write(path2write)

                    #doc
                    DocxListObj.add_run("\n" + path2write + "\n").font.size = Pt(FontSizeMain)

                    #csv
                    # header = ['index', 'full_path', 'path_no_file_name', 'file_name', 'label']
                    path_no_file_name, file_name = os.path.split(path2write)
                    file_label = get_obj_label_from_file_name(file_name=file_name)

                    row_csv = [str(count), path2write , path_no_file_name , file_name , file_label ]
                    # write the header
                    writer.writerow(row_csv)



                    if count >= self.limit:
                        list_file_txt.close()
                        list_3Dmodels_txt_count += 1
                        list_file_txt = open(f"{self.list_3Dmodels_txt}_{list_3Dmodels_txt_count}.txt", 'w')
                        count = 0

        list_file_txt.close()
        self.path_list_doc = path.normpath(path.abspath(path.join(self.out_model3D_list , self.list_3Dmodels_docx + ".docx" )))
        list_file_doc.save(self.path_list_doc)
        list_file_csv.close()
        print('Created list of 3D models with labels')

if __name__ == "__main__":
    # lis = ['test', 'train']
    # for name in lis:
    ROOT_DIR = os.path.dirname(os.path.abspath(
        "top_level_file.txt"))  # Get project root https://www.adamsmith.haus/python/answers/how-to-get-the-path-of-the-root-project-structure-in-python

    models_3D_list(model3D_dir= ROOT_DIR, out_model3D_list=ROOT_DIR)  # limit=2000